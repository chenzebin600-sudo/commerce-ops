from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "showcase"
OUTPUT_PATH = OUTPUT_DIR / "Commerce-Ops-AI-Project-Portfolio.docx"

SCREENSHOTS = {
    "growth": ROOT / "docs" / "screenshots" / "growth-radar-v2-mainline-desktop.png",
    "media": ROOT / "docs" / "screenshots" / "mabang-sku-image-real-session-desktop.png",
    "listing": ROOT / "docs" / "screenshots" / "product-listing-ai-positioning-title.png",
}

NAVY = "102A43"
TEAL = "0F766E"
BLUE = "2563EB"
SKY = "EAF3F8"
PALE_TEAL = "E9F7F4"
GOLD = "B7791F"
PALE_GOLD = "FFF7E6"
RED = "B42318"
PALE_RED = "FDECEC"
INK = "243B53"
MUTED = "627D98"
LIGHT = "F4F7FA"
BORDER = "D8E2EA"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Iterable[int], indent=TABLE_INDENT_DXA):
    widths = list(widths)
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border_bottom(paragraph, color=BORDER, size=8, space=3):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_keep(paragraph, keep_next=False, keep_lines=True):
    paragraph.paragraph_format.keep_together = keep_lines
    paragraph.paragraph_format.keep_with_next = keep_next


def add_para(
    doc,
    text="",
    *,
    size=11,
    color=INK,
    bold=False,
    italic=False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before=0,
    after=6,
    line=1.1,
    keep=False,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    if keep:
        set_keep(p, keep_lines=True)
    return p


def add_mixed_para(doc, parts, *, after=6, before=0, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for part in parts:
        run = p.add_run(part.get("text", ""))
        set_run_font(
            run,
            size=part.get("size", 11),
            color=part.get("color", INK),
            bold=part.get("bold"),
            italic=part.get("italic"),
        )
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_bullet(doc, text, *, level=0, bold_prefix=None):
    style = "AI Bullet" if level == 0 else "AI Bullet 2"
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, size=10.7, bold=True, color=INK)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=10.7, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.7, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="AI Number")
    run = p.add_run(text)
    set_run_font(run, size=10.7, color=INK)
    return p


def add_callout(doc, label, text, *, fill=PALE_TEAL, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=3)
    return table


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [CONTENT_WIDTH_DXA // len(metrics)] * len(metrics)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths, indent=0)
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT if idx % 2 == 0 else PALE_TEAL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=18, color=TEAL if idx % 2 else NAVY, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_run_font(r2, size=8.6, color=MUTED, bold=True)
    return table


def add_table(doc, headers, rows, widths, *, header_fill=NAVY, first_col_bold=False, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=9, color=WHITE, bold=True)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        tr_pr = table.rows[-1]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, LIGHT)
        for idx, text in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if idx > 0 and len(str(text)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(text))
            set_run_font(run, size=font_size, color=INK, bold=first_col_bold and idx == 0)
    add_para(doc, "", after=2)
    return table


def add_screenshot(doc, path, caption, *, width=6.25):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    image = run.add_picture(str(path), width=Inches(width))
    image._inline.docPr.set("title", caption)
    image._inline.docPr.set("descr", caption)
    cap = add_para(
        doc,
        caption,
        size=8.5,
        color=MUTED,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.0,
    )
    set_keep(cap, keep_lines=True)


def add_page_break(doc):
    doc.add_page_break()


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        1: (16, TEAL, 16, 8),
        2: (13, TEAL, 12, 6),
        3: (11.5, NAVY, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles.add_style("AI Bullet", 1)
    bullet.base_style = styles["List Bullet"]
    bullet.font.name = "Microsoft YaHei"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(10.7)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.1

    bullet2 = styles.add_style("AI Bullet 2", 1)
    bullet2.base_style = styles["List Bullet 2"]
    bullet2.font.name = "Microsoft YaHei"
    bullet2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet2.font.size = Pt(10.4)
    bullet2.paragraph_format.left_indent = Inches(0.75)
    bullet2.paragraph_format.first_line_indent = Inches(-0.2)
    bullet2.paragraph_format.space_after = Pt(3)

    number = styles.add_style("AI Number", 1)
    number.base_style = styles["List Number"]
    number.font.name = "Microsoft YaHei"
    number._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    number.font.size = Pt(10.7)
    number.paragraph_format.left_indent = Inches(0.5)
    number.paragraph_format.first_line_indent = Inches(-0.25)
    number.paragraph_format.space_after = Pt(5)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.38)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.add_run("COMMERCE OPS  |  AI PROJECT PORTFOLIO")
    set_run_font(p.runs[0], size=8.5, color=MUTED, bold=True)
    set_paragraph_border_bottom(p, color=BORDER, size=6, space=4)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AI 项目作品集  |  ")
    set_run_font(run, size=8.2, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def cover_page(doc):
    add_para(doc, "AI PROJECT PORTFOLIO", size=10, color=TEAL, bold=True, after=12)
    title = add_para(doc, "Commerce Ops", size=31, color=NAVY, bold=True, after=3, line=1.0)
    set_keep(title, keep_next=True)
    add_para(
        doc,
        "从业务问题到可验证系统：一个 AI 驱动的跨境电商运营平台",
        size=16,
        color=TEAL,
        bold=True,
        after=16,
        line=1.15,
    )
    add_para(
        doc,
        "AI 产品负责人 / 业务架构发起人 / 人机协同研发实践",
        size=10.5,
        color=MUTED,
        bold=True,
        after=24,
    )
    quote = add_para(
        doc,
        "“我不是让 AI 替我做决定，而是用 AI 把业务经验变成可测试、可审计、可演进的软件系统。”",
        size=14,
        color=INK,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=6,
        after=22,
        line=1.3,
    )
    set_paragraph_border_bottom(quote, color=TEAL, size=10, space=8)
    add_metric_strip(
        doc,
        [
            ("18,342", "活跃产品 / SKU"),
            ("6,583", "图片资产"),
            ("718/718", "全量测试通过"),
            ("41,156", "审计事件"),
        ],
    )
    add_para(doc, "", after=10)
    add_screenshot(
        doc,
        SCREENSHOTS["growth"],
        "Commerce Ops 统一工作台中的 Growth Radar V2.2 超级店长运营助手（演示数据已脱敏）",
        width=5.85,
    )
    add_para(
        doc,
        "展示版说明：真实账号、店铺、商品、文件路径和密钥均不进入本作品集；规模数据来自 2026-07-28 正式库只读审计。",
        size=8.2,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
        line=1.0,
    )
    add_page_break(doc)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    props = doc.core_properties
    props.title = "Commerce Ops AI 项目作品集"
    props.subject = "AI 驱动的跨境电商运营平台案例"
    props.keywords = "AI, Commerce Ops, 电商运营, Growth Radar, 马帮, React, Node.js, Python"
    props.comments = "对外展示版；敏感信息已省略。"

    cover_page(doc)

    add_heading(doc, "01  AI 项目摘要", 1)
    add_callout(
        doc,
        "一句话定位",
        "Commerce Ops 是一个连接马帮 ERP、产品资料、图片资产、经营分析和 Listing 执行的跨境电商运营控制台。",
    )
    add_para(
        doc,
        "这个项目不是单一网页，也不是一次性脚本。它从真实业务问题出发，逐步形成了数据采集、事实治理、确定性分析、运营任务、图片资产和平台执行的闭环。",
        after=8,
        line=1.2,
    )
    add_heading(doc, "我的项目角色", 2)
    add_table(
        doc,
        ["角色维度", "我承担的工作", "AI 协作方式"],
        [
            ("业务负责人", "定义电商运营痛点、指标口径和优先级", "AI 帮助澄清需求、识别歧义、形成产品合同"),
            ("产品负责人", "把目标拆成可验收节点和页面流程", "AI 生成 PRD、线框、任务包和验收清单"),
            ("架构发起人", "决定数据真源、模块边界和高风险门禁", "AI 阅读代码、比较方案、输出架构审计"),
            ("研发协调者", "推进前后端、数据库、测试和集成", "Codex 执行实现、测试、构建和浏览器验证"),
            ("质量负责人", "保护正式数据库并确认发布条件", "隔离迁移、全量测试、哈希校验和视觉 QA"),
        ],
        [1650, 3750, 3960],
        first_col_bold=True,
        font_size=9,
    )
    add_heading(doc, "这项 AI 实践证明了什么", 2)
    for item in [
        "AI 可以参与长期软件项目，而不只是回答问题或生成单段代码。",
        "业务人员可以通过数据合同、节点门禁和测试结果，稳定管理 AI 研发过程。",
        "人负责目标、业务真相和高风险批准；AI 负责阅读、实现、验证和文档化。",
        "真正有价值的“AI 成功”不是提示词数量，而是可运行、可审计、可迭代的系统结果。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "02  业务问题与闭环", 1)
    add_para(
        doc,
        "跨境电商运营通常同时面对多国家、多平台、多店铺和大量 SKU。一个店长可能负责数十家店铺，订单、库存、图片和 Listing 分散在不同系统中，异常与机会很容易被遗漏。",
        after=8,
    )
    add_heading(doc, "核心业务挑战", 2)
    for item in [
        "订单和库存每天变化，但数据通常停留在 Excel 或来源系统页面。",
        "市场已经验证的货盘，与自己店铺的销售表现缺少横向对比。",
        "图片散落在马帮、产品包和人工资料中，重复保存且难以回到产品 SKU。",
        "Listing 草稿、图片、平台属性和发布结果缺少一致的工作流。",
        "大量任务靠人工记忆，无法按优先级、证据和负责人持续跟踪。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "系统闭环", 2)
    add_table(
        doc,
        ["1 数据获取", "2 事实治理", "3 分析判断", "4 运营任务", "5 执行回读"],
        [
            (
                "马帮订单、库存、图片、在线商品",
                "来源批次、标准化、映射和质量检查",
                "确定性指标、机会和风险",
                "负责人、优先级、证据和动作",
                "产品、图片、Listing、平台结果",
            )
        ],
        [1872, 1872, 1872, 1872, 1872],
        header_fill=TEAL,
        font_size=8.7,
    )
    add_callout(
        doc,
        "关键原则",
        "所有推荐都必须回答“为什么、依据是什么、建议做什么”；AI 不得绕过规则和人工确认直接执行经营动作。",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "03  系统能力地图", 1)
    add_table(
        doc,
        ["模块", "已经实现的能力", "成熟度"],
        [
            ("产品中心", "产品包导入、SKU 主数据、包装/成本、字段覆盖、软删除、正式图片、AI 内容、Listing 草稿", "正式启用"),
            ("马帮数据中心", "账号管理、订单/库存采集、定时导出、Excel 证据、自动入库、失败重试、钉钉通知", "正式启用"),
            ("Growth Radar", "数据范围治理、国家类目机会、SKU 指标、店铺诊断、确定性信号、店长任务", "A2 正式；V2.2 待正式迁移"),
            ("图片素材", "SKU 图片发现、下载校验、SHA 去重、全量分段、断点恢复、产品关联和主图保护", "正式启用"),
            ("Listing 工作台", "草稿、AI 内容、图片编排、发布检查、马帮在线商品查询、安全批改、发布和回读", "主库 + 隔离侧车"),
            ("任务系统", "定时任务、图片任务、运营任务、发布任务、AI 图片任务、文件审核任务", "多套实现，待统一"),
            ("竞品与广告", "链接维度竞品分析、关键词发现、主图/广告素材分析", "可用能力"),
            ("文件与审计", "文件登记、扫描、隔离、恢复、人工审核、HTTP 与业务操作审计", "正式启用"),
            ("AI Gateway", "DeepSeek 接入、结构化输出、内容候选、确定性回退、审计和密钥隔离", "正式启用"),
        ],
        [1500, 5960, 1900],
        first_col_bold=True,
        font_size=8.7,
    )

    add_heading(doc, "04  功能全景：产品与数据", 1)
    add_heading(doc, "产品中心", 2)
    for item in [
        "批量导入产品包 Excel，保留来源文件、原始行、问题和字段变更证据。",
        "管理类目、型号、SKU、国家、仓库、包装、成本和库存投影。",
        "支持字段人工覆盖、修改历史、软删除和恢复。",
        "管理正式产品图片，避免参考素材自动覆盖人工主图。",
        "生成并保存 AI 标题、副标题、描述、卖点、使用场景和图片候选。",
        "创建平台/国家/店铺维度的 Listing 草稿并执行发布前检查。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "马帮数据中心", 2)
    for item in [
        "集中管理加密马帮账号并执行登录验证。",
        "手工或定时获取订单、库存数据，支持日期、仓库和业务筛选。",
        "保存 Excel 文件及 SHA-256，形成可追溯来源证据。",
        "把订单和库存写入统一事实层，而不是只停留在页面或 Excel。",
        "支持调度租约、漏跑处理、失败重试、文件保留和钉钉通知。",
        "保留 WPS 桌面同步助手，兼容人工工作流。",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "数据治理",
        "订单和库存以来源批次、原始证据、标准事实和业务投影分层；Excel 是证据，不是分析数据库。",
    )

    add_heading(doc, "05  功能全景：增长与运营", 1)
    add_heading(doc, "Growth Radar V2.2：超级店长运营助手", 2)
    for item in [
        "今日作战台：每位店长最多展示 10 项最重要任务。",
        "店铺战场：按店长、店铺、国家和平台查看健康状态、异常和机会。",
        "国家 × 类目机会地图：发现值得投入的国家和类目方向。",
        "产品雷达：识别明星、增长、衰退、蓝海和跨国候选 SKU。",
        "货盘表现 vs 我方表现：用四象限解释最大机会、核心产品、优势产品和低优先产品。",
        "店铺缺口诊断：识别高表现货盘中店铺承接不足的 SKU。",
        "SKU 详情：展示预测日销量、我方销量、库存、规则、公式和证据。",
        "运营任务生命周期：确认、处理中、观察、解决、阻塞、忽略和重新打开。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "确定性分析规则", 2)
    add_table(
        doc,
        ["规则", "输入", "输出"],
        [
            ("趋势", "当前 7 天 vs 前 7 天有效订单销量", "增长、下降、数据不足"),
            ("货盘验证", "国家、类目、SKU、预测日销量、P80", "高表现货盘排名"),
            ("低承接", "货盘表现 vs 我方有效销量，阈值 10%", "优先发力方向"),
            ("库存风险", "国家 + 仓库 + SKU 的可售天数", "缺货、预警、正常"),
            ("滞销", "活跃度、库存、60/90/180 天阈值", "关注、风险、严重"),
            ("新品", "是否新款 + 90 天观察周期", "新品机会与覆盖情况"),
        ],
        [1700, 4400, 3260],
        first_col_bold=True,
        font_size=9,
    )
    add_screenshot(
        doc,
        SCREENSHOTS["growth"],
        "Growth Radar V2.2：任务优先、证据可解释、正式数据门禁可见",
        width=6.2,
    )

    add_page_break(doc)
    add_heading(doc, "06  功能全景：图片资产", 1)
    for item in [
        "从马帮库存页发现 SKU 与图片，支持真实账号后台登录。",
        "下载前校验域名、响应类型、文件签名、大小和尺寸。",
        "下载过程中计算 SHA-256，相同内容只保存一次物理文件。",
        "支持单批、缺失补采、失败重试和全量分段同步。",
        "保存分页检查点，支持暂停、恢复和进程重启后的安全恢复。",
        "图片可关联多个产品、国家和 SKU，但保留各自来源证据。",
        "马帮图片默认只是参考素材，不能自动设置为正式主图。",
        "用户可以确认图库、确认主图、拒绝关联或后续生成白底图候选。",
    ]:
        add_bullet(doc, item)
    add_metric_strip(
        doc,
        [
            ("6,583", "去重图片资产"),
            ("33,759", "产品素材关联"),
            ("21,460", "库存快照"),
            ("2", "受管马帮账号"),
        ],
    )
    add_para(doc, "", after=8)
    add_screenshot(
        doc,
        SCREENSHOTS["media"],
        "COM-015 马帮 SKU 图片采集：账号、批次、断点、去重与产品关联（敏感信息已脱敏）",
        width=5.85,
    )

    add_heading(doc, "07  功能全景：Listing 与执行", 1)
    add_heading(doc, "产品中心 Listing 工作台", 2)
    for item in [
        "按平台、国家、店铺和类目建立草稿。",
        "编辑标题、副标题、描述、卖点、场景、属性、变体、价格、库存和物流。",
        "编排图片顺序、主图和视频，并检查必填项、阻断项和完整度。",
        "基于产品事实生成 AI 内容候选，记录采用、人工修改和历史版本。",
        "发布前保持人工确认，避免 AI 或后台任务直接发布。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "马帮在线商品与刊登侧车", 2)
    for item in [
        "查询 Lazada、Shopee、TikTok Shop 已授权在线商品。",
        "批量修改前生成预览，并执行明确确认、过期检查和串行执行。",
        "支持复制在线商品、创建草稿、校验、确认发布和状态轮询。",
        "发布完成后回读平台结果；侧车数据库和审计日志与正式库隔离。",
        "浏览器只访问 Commerce Ops 主服务，内部代理和令牌不会暴露给前端。",
    ]:
        add_bullet(doc, item)
    add_screenshot(
        doc,
        SCREENSHOTS["listing"],
        "产品中心 Listing 工作台：AI 候选、人工采用、平台字段与发布前检查",
        width=6.25,
    )

    add_page_break(doc)
    add_heading(doc, "08  AI 能力与人机协作边界", 1)
    add_heading(doc, "项目中的 AI 应用", 2)
    add_table(
        doc,
        ["AI 能力", "实际用途", "控制机制"],
        [
            ("需求与架构协作", "把业务描述转成 PRD、指标合同、数据模型和实施节点", "人确认口径与边界"),
            ("代码协作", "代码阅读、跨模块集成、测试补充、问题定位和重构建议", "Git 边界与测试门禁"),
            ("内容生成", "标题、描述、卖点、场景、定位和多语言候选", "候选制、人工采用、版本历史"),
            ("图片辅助", "生成计划、白底图/场景图候选和历史版本", "不自动覆盖正式图片"),
            ("结构化指令", "把自然语言转换为受控 Listing 操作", "Schema 校验 + 确定性回退"),
            ("经营解释", "基于确定性指标组织原因、证据和建议动作", "不允许 AI 直接评分或执行"),
        ],
        [1800, 4300, 3260],
        first_col_bold=True,
        font_size=9,
    )
    add_heading(doc, "人机分工", 2)
    add_table(
        doc,
        ["由人负责", "由 AI / Codex 协助"],
        [
            ("业务真相、指标口径、优先级", "代码库阅读、差异分析和方案生成"),
            ("正式数据库、migration、发布和全量任务批准", "隔离实现、测试、构建和验证"),
            ("经营动作和主图确认", "候选内容、证据整理和操作建议"),
            ("风险接受与最终决策", "问题分类、复盘报告和架构审计"),
        ],
        [4680, 4680],
        header_fill=TEAL,
        font_size=9.4,
    )
    add_callout(
        doc,
        "AI 治理原则",
        "AI 可以提出建议、生成候选和执行已授权工程任务，但不能自行改变业务口径、正式数据、迁移历史或高风险运营动作。",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "09  技术栈", 1)
    add_table(
        doc,
        ["层次", "技术与工程实践"],
        [
            ("前端", "React 19、TypeScript、Vite 8、Ant Design、Tailwind CSS 4、ECharts 6、Fluent UI、Lucide、原生 ES Modules、Shadow DOM React Island"),
            ("后端", "Node.js ESM、原生 HTTP 服务、REST API、模块化服务层、Python worker/sidecar、子进程管理、固定回环代理"),
            ("数据", "SQLite、SQL migrations、Repository/Provider 模式、事务、幂等键、来源批次、原始层/事实层/投影层、PostgreSQL readiness"),
            ("集成", "马帮 ERP、WPS/Excel、Chrome/CDP、钉钉机器人、Lazada、Shopee、TikTok Shop、DeepSeek"),
            ("文件与图片", "本地受管存储、相对 storage key、SHA-256 去重、MIME/签名/尺寸校验、临时文件、隔离与恢复"),
            ("安全", "账号加密、访问令牌、内部侧车令牌、loopback 限制、网络 allowlist、PII 过滤、Excel 公式注入防护、操作审计"),
            ("测试与质量", "Node test runner、Python unittest、TypeScript check、Vite build、浏览器/移动端验证、隔离 migration、Doctor、正式库哈希保护"),
            ("AI 工程", "AI Gateway、DeepSeek provider、结构化输出、Schema 校验、确定性回退、候选采用、版本和审计"),
        ],
        [1500, 7860],
        first_col_bold=True,
        font_size=9,
    )

    add_heading(doc, "10  数据架构与系统设计", 1)
    add_heading(doc, "数据分层", 2)
    add_table(
        doc,
        ["数据层", "示例", "职责"],
        [
            ("来源证据", "Excel、raw rows、source batch、文件哈希", "证明数据从哪里来、何时获取、范围是什么"),
            ("业务事实", "订单头/行、库存快照、产品 SKU", "提供可重算、可查询的真实数据"),
            ("业务投影", "产品库存、SKU 覆盖、图片关联", "为具体页面和业务流程提供视图"),
            ("分析结果", "analysis run、指标、信号", "保存规则版本、公式、输入和证据"),
            ("运营状态", "任务、负责人、优先级、事件", "保存人的处理过程，不被每日重算覆盖"),
            ("执行结果", "Listing 发布记录、平台回读、审计", "确认系统做过什么、结果如何"),
        ],
        [1500, 3280, 4580],
        first_col_bold=True,
        font_size=9,
    )
    add_heading(doc, "架构模式", 2)
    for item in [
        "模块化主服务：产品、事实、分析、媒体、Listing 和任务各自拥有领域服务。",
        "受管侧车：Python 只处理外部协议和桌面集成，主服务负责认证、代理和生命周期。",
        "渐进式前端：在统一工作台中通过 React Island 逐步替换旧页面，不使用 iframe。",
        "Fail-closed：数据、迁移或规则门禁未满足时，页面展示阻断原因，不生成虚假建议。",
        "正式库保护：开发和验收使用临时 SQLite 或正式库复制件，最终用 SHA-256 验证未漂移。",
    ]:
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "11  可量化工程成果", 1)
    add_metric_strip(
        doc,
        [
            ("18,347", "产品 / SKU 总量"),
            ("21,714", "产品包原始行"),
            ("21,460", "库存快照"),
            ("2,726", "订单商品行"),
        ],
    )
    add_para(doc, "", after=8)
    add_metric_strip(
        doc,
        [
            ("6,583", "图片资产"),
            ("33,759", "图片关联"),
            ("41,156", "审计事件"),
            ("61", "正式数据库表"),
        ],
    )
    add_heading(doc, "质量证据", 2)
    add_table(
        doc,
        ["验证项", "结果", "说明"],
        [
            ("主项目全量测试", "718 / 718", "0 失败，覆盖产品、Growth、图片、调度、文件和集成"),
            ("Mabang-getdata Python", "58 / 58", "保留来源项目原能力并补充内部认证测试"),
            ("马帮刊登集成", "9 / 9", "包含真实 Python 子进程、内部令牌和隔离存储"),
            ("原刊登看板", "3 / 3", "构建和渲染测试通过"),
            ("前端结构检查", "470 / 219", "470 个唯一元素 ID、219 个静态绑定"),
            ("数据库", "PASS", "integrity_check=ok，外键异常 0"),
            ("视觉验证", "PASS", "桌面端与 430px 移动端通过，无横向溢出和控制台错误"),
            ("正式库保护", "PASS", "SQLite、WAL、SHM 哈希在整合验收前后保持一致"),
        ],
        [2300, 1700, 5360],
        first_col_bold=True,
        font_size=9,
    )
    add_callout(
        doc,
        "可信度说明",
        "这些数字代表系统规模和工程验证，不等同于商业收入或生产收益；作品集不虚构 GMV、转化率或节省工时。",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "12  我形成的可迁移能力", 1)
    add_table(
        doc,
        ["能力", "可证明的表现"],
        [
            ("AI 产品能力", "把模糊业务问题转化为指标、页面、数据合同和验收标准"),
            ("业务架构能力", "识别数据真源、投影、候选能力和正式启用边界"),
            ("AI 协同研发", "用阶段任务包管理 Codex，实现、测试、复盘并持续交接"),
            ("数据治理", "建立来源批次、原始证据、事实层、规则版本和审计链"),
            ("风险控制", "对 migration、正式数据库、批量采集和平台发布设置人工门禁"),
            ("全栈沟通", "能够在前端、后端、数据库、Python 集成和业务运营之间做取舍"),
            ("质量意识", "要求全量测试、Build、Doctor、浏览器、移动端和哈希共同证明结果"),
            ("长期演进", "能从单点工具逐步演进到模块化运营平台，并主动识别技术债"),
        ],
        [2100, 7260],
        first_col_bold=True,
        font_size=9.4,
    )

    add_heading(doc, "13  当前边界与下一步", 1)
    add_heading(doc, "已经完成", 2)
    for item in [
        "产品中心、马帮订单/库存事实层、图片素材和主工作台已形成可用基础。",
        "Growth Radar V2.2 的设计、运行时、前端和候选迁移已完成验证。",
        "Mabang-getdata 已整合到主项目，来源功能和独立运行能力得到保留。",
        "系统架构、数据真源、重复能力和技术债已经形成正式审计文档。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "仍需正式批准或收敛", 2)
    for item in [
        "Growth Radar 019/020/021 正式迁移与首轮正式分析。",
        "主库 Listing 草稿与 Python 侧车草稿的单一真源。",
        "马帮账号/会话代理和跨模块任务基础设施。",
        "候选 migration 与正式启动之间的显式部署门禁。",
        "前端共享运行时、巨型文件拆分和统一可观测性。",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "架构取舍",
        "当前不急于微服务化，也不急于整体迁移 PostgreSQL 或 MinIO；先收敛数据真源、账号、任务和部署边界。",
    )

    add_page_break(doc)
    add_heading(doc, "14  可直接放入简历的项目描述", 1)
    add_callout(
        doc,
        "项目名称",
        "Commerce Ops - AI 驱动的跨境电商运营平台",
    )
    add_mixed_para(
        doc,
        [
            {"text": "项目角色：", "bold": True, "color": NAVY},
            {"text": "AI 产品负责人 / 业务架构发起人 / 人机协同研发实践"},
        ],
        after=5,
    )
    add_mixed_para(
        doc,
        [
            {"text": "项目简介：", "bold": True, "color": NAVY},
            {
                "text": "围绕多国家、多平台、多店铺运营场景，使用 ChatGPT 与 Codex 协同完成从需求澄清、数据合同、系统架构到前后端实现、数据库迁移演练和全量验收的长期项目。系统连接马帮 ERP，覆盖产品、订单、库存、图片、增长分析、运营任务和 Listing 执行。"
            },
        ],
        after=8,
        line=1.2,
    )
    add_heading(doc, "简历成果要点", 2)
    for item in [
        "设计并落地订单/库存事实层，使马帮定时采集、Excel 证据和 Growth Radar 使用同一数据来源。",
        "构建 Growth Radar 超级店长助手，以国家、类目、店铺、店长和 SKU 为维度生成可解释的确定性机会与任务。",
        "实现马帮 SKU 图片全量采集、SHA-256 去重、断点恢复和产品关联，正式库积累 6,583 个资产与 33,759 条关联。",
        "整合马帮在线商品查询、批量安全修改和刊登发布能力，并通过主服务代理隔离账号与侧车数据。",
        "建立正式数据库保护机制，全量测试达到 718/718，通过桌面、移动端、Build、Doctor 和哈希一致性验证。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "技术栈简写", 2)
    add_para(
        doc,
        "React 19 / TypeScript / Vite / Ant Design / Tailwind CSS / ECharts / Fluent UI / Node.js / Python / SQLite / REST / Chrome-CDP / WPS-Excel / DeepSeek / SHA-256 / Playwright-style browser validation",
        size=10,
        color=INK,
        after=10,
        line=1.25,
    )
    add_heading(doc, "面试讲述提纲", 2)
    for item in [
        "先讲真实业务难点：一个运营需要管理大量店铺和 SKU，信息分散且无法持续关注异常。",
        "再讲 AI 协作方式：人冻结业务口径和风险边界，AI 执行阅读、实现、测试和复盘。",
        "再讲最难的技术点：数据真源、迁移门禁、图片去重、任务生命周期和侧车隔离。",
        "用量化证据收尾：系统规模、测试结果、正式库保护和跨端验证。",
        "最后说明边界：哪些已经正式启用，哪些仍是候选，为什么没有为了展示而冒险上线。",
    ]:
        add_number(doc, item)

    add_heading(doc, "结语", 1)
    add_para(
        doc,
        "这个项目展示的不是“AI 替我写了多少代码”，而是我如何把业务知识、AI 协作、工程纪律和风险控制组合成一个持续演进的软件系统。",
        size=13,
        color=TEAL,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=8,
        after=8,
        line=1.3,
    )
    add_para(
        doc,
        "Commerce Ops 证明：业务人员可以不止使用 AI，还可以管理 AI、验证 AI，并把 AI 变成可靠的交付能力。",
        size=11,
        color=MUTED,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
    )

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
